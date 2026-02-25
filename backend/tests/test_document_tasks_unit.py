"""
Unit tests for document_tasks Celery tasks.
Tests _hash_content, _extract_docx_text, and process_document logic.
"""

import hashlib
from contextlib import asynccontextmanager
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from app.tasks.document_tasks import _hash_content


def _mock_session_ctx(mock_session):
    @asynccontextmanager
    async def _ctx():
        yield mock_session

    return _ctx


class TestHashContent:
    def test_basic_hash(self):
        text = "hello world"
        expected = hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()
        assert _hash_content(text) == expected

    def test_empty_string(self):
        result = _hash_content("")
        assert isinstance(result, str)
        assert len(result) == 64

    def test_unicode_content(self):
        text = "cafe\u0301 \u2603 \U0001f600"
        result = _hash_content(text)
        assert isinstance(result, str)
        assert len(result) == 64


class TestExtractDocxText:
    def test_extract_from_valid_docx(self):
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        import io

        doc = DocxDocument()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        buffer = io.BytesIO()
        doc.save(buffer)

        from app.tasks.document_tasks import _extract_docx_text

        result = _extract_docx_text(buffer.getvalue())
        assert "First paragraph" in result
        assert "Second paragraph" in result

    def test_empty_docx_returns_empty(self):
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        import io

        doc = DocxDocument()
        buffer = io.BytesIO()
        doc.save(buffer)

        from app.tasks.document_tasks import _extract_docx_text

        result = _extract_docx_text(buffer.getvalue())
        assert result == ""


class TestProcessDocument:
    def test_document_not_found(self):
        mock_session = AsyncMock()
        mock_result = MagicMock()
        mock_result.scalar_one_or_none.return_value = None
        mock_session.execute.return_value = mock_result

        with patch(
            "app.tasks.document_tasks.get_celery_session_context",
            _mock_session_ctx(mock_session),
        ):
            from app.tasks.document_tasks import process_document

            result = process_document.apply(kwargs={"document_id": 9999})
            assert result.result["status"] == "error"
            assert "not found" in result.result["error"]

    def test_file_not_found(self):
        mock_doc = MagicMock()
        mock_doc.file_path = "/nonexistent/path/to/doc.pdf"
        mock_doc.processing_status = None
        mock_doc.processing_error = None
        mock_doc.id = 1

        mock_session = AsyncMock()
        mock_result = MagicMock()
        mock_result.scalar_one_or_none.return_value = mock_doc
        mock_session.execute.return_value = mock_result

        with patch(
            "app.tasks.document_tasks.get_celery_session_context",
            _mock_session_ctx(mock_session),
        ):
            from app.tasks.document_tasks import process_document

            result = process_document.apply(kwargs={"document_id": 1})
            assert result.result["status"] == "error"

    def test_unsupported_file_type(self, tmp_path):
        test_file = tmp_path / "test.xyz"
        test_file.write_bytes(b"some content")

        mock_doc = MagicMock()
        mock_doc.file_path = str(test_file)
        mock_doc.mime_type = "application/unknown"
        mock_doc.original_filename = "test.xyz"
        mock_doc.processing_status = None
        mock_doc.processing_error = None
        mock_doc.id = 1

        mock_session = AsyncMock()
        mock_result = MagicMock()
        mock_result.scalar_one_or_none.return_value = mock_doc
        mock_session.execute.return_value = mock_result

        with patch(
            "app.tasks.document_tasks.get_celery_session_context",
            _mock_session_ctx(mock_session),
        ):
            from app.tasks.document_tasks import process_document

            result = process_document.apply(kwargs={"document_id": 1})
            assert result.result["status"] == "error"
            assert "Unsupported" in result.result["error"]

    def test_text_file_processing(self, tmp_path):
        test_file = tmp_path / "test.txt"
        test_file.write_text("Hello world test content")

        mock_doc = MagicMock()
        mock_doc.file_path = str(test_file)
        mock_doc.mime_type = "text/plain"
        mock_doc.original_filename = "test.txt"
        mock_doc.processing_status = None
        mock_doc.processing_error = None
        mock_doc.id = 1
        mock_doc.user_id = 1
        mock_doc.title = "Test"
        mock_doc.document_type = MagicMock(value="capability_statement")
        mock_doc.description = None
        mock_doc.full_text = None

        mock_session = AsyncMock()
        mock_result = MagicMock()
        mock_result.scalar_one_or_none.return_value = mock_doc
        mock_session.execute.return_value = mock_result

        with (
            patch(
                "app.tasks.document_tasks.get_celery_session_context",
                _mock_session_ctx(mock_session),
            ),
            patch("app.tasks.document_tasks.index_entity", new_callable=AsyncMock),
        ):
            from app.tasks.document_tasks import process_document

            result = process_document.apply(kwargs={"document_id": 1})
            assert result.result["status"] == "completed"
            assert result.result["page_count"] == 1
            assert result.result["chunks"] == 1

    def test_legacy_doc_raises(self, tmp_path):
        test_file = tmp_path / "test.doc"
        test_file.write_bytes(b"\xd0\xcf\x11\xe0legacy doc content")

        mock_doc = MagicMock()
        mock_doc.file_path = str(test_file)
        mock_doc.mime_type = "application/msword"
        mock_doc.original_filename = "test.doc"
        mock_doc.processing_status = None
        mock_doc.processing_error = None
        mock_doc.id = 1

        mock_session = AsyncMock()
        mock_result = MagicMock()
        mock_result.scalar_one_or_none.return_value = mock_doc
        mock_session.execute.return_value = mock_result

        with patch(
            "app.tasks.document_tasks.get_celery_session_context",
            _mock_session_ctx(mock_session),
        ):
            from app.tasks.document_tasks import process_document

            result = process_document.apply(kwargs={"document_id": 1})
            assert result.result["status"] == "error"
            assert ".doc" in result.result["error"]
