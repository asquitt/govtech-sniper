"""
RFP Sniper - Export Routes
===========================
Export proposals to DOCX and PDF formats.
"""

import html
import io
import json
import re
import zipfile
from collections import defaultdict
from datetime import datetime
from typing import Any

import structlog
from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlmodel import select

from app.api.deps import (
    STEP_UP_REQUIRED_HEADER,
    UserAuth,
    get_current_user,
    get_step_up_code,
    get_user_org_security_policy,
    get_user_policy_role,
    verify_step_up_code,
)
from app.database import get_session
from app.models.capture import BidScorecard
from app.models.knowledge_base import DocumentChunk, KnowledgeBaseDocument
from app.models.proposal import Proposal, ProposalSection, SectionEvidence
from app.models.review import ProposalReview, ReviewComment
from app.models.rfp import RFP, ComplianceMatrix
from app.services.audit_service import log_audit_event
from app.services.policy_engine import PolicyAction, PolicyDecision, PolicyResult, evaluate

logger = structlog.get_logger(__name__)

router = APIRouter(prefix="/export", tags=["Export"])


# =============================================================================
# HTML-to-DOCX helpers
# =============================================================================

_TAG_RE = re.compile(r"<[^>]+>")


async def _enforce_export_policy(
    *,
    policy: PolicyResult,
    current_user: UserAuth,
    session: AsyncSession,
    request: Request,
    step_up_code: str | None = None,
) -> None:
    if policy.decision == PolicyDecision.DENY:
        await session.commit()
        raise HTTPException(status_code=403, detail=policy.reason)
    if policy.decision != PolicyDecision.STEP_UP:
        return

    org_policy = await get_user_org_security_policy(current_user.id, session)
    if not org_policy.get("require_step_up_for_sensitive_exports", True):
        return

    supplied_code = get_step_up_code(request, step_up_code)
    if await verify_step_up_code(current_user.id, session, supplied_code):
        await log_audit_event(
            session,
            user_id=current_user.id,
            entity_type="security",
            entity_id=current_user.id,
            action="security.step_up.challenge_succeeded",
            metadata={
                "channel": "export",
                **policy.to_audit_dict(),
            },
        )
        await session.commit()
        return

    await log_audit_event(
        session,
        user_id=current_user.id,
        entity_type="security",
        entity_id=current_user.id,
        action="security.step_up.challenge_failed",
        metadata={
            "channel": "export",
            **policy.to_audit_dict(),
        },
    )
    await session.commit()
    raise HTTPException(
        status_code=403,
        detail=policy.reason,
        headers={STEP_UP_REQUIRED_HEADER: "true"},
    )


def _strip_tags(content: str) -> str:
    """Remove all HTML tags, returning plain text."""
    return html.unescape(_TAG_RE.sub("", content)).strip()


def _has_html(text: str) -> bool:
    return bool(re.search(r"<[a-z][\s\S]*>", text, re.IGNORECASE))


def _render_html_to_docx(doc: "Document", html: str) -> None:  # type: ignore[name-defined]
    """Convert HTML content from TipTap into python-docx paragraphs.

    Handles <h1>-<h4>, <p>, <blockquote>, <ul>/<ol> <li>, and strips
    inline tags (<strong>, <em>) into plain text. For content that isn't
    HTML (legacy plain-text), falls back to newline splitting.
    """
    if not _has_html(html):
        for para in html.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        return

    # Expand <ul>/<ol> blocks into individual <li> tags for matching
    def _expand_lists(h: str) -> str:
        return re.sub(
            r"<(ul|ol)[^>]*>(.*?)</\1>",
            lambda m: m.group(2),
            h,
            flags=re.DOTALL,
        )

    expanded = _expand_lists(html)

    # Match top-level block elements in order
    block_pattern = re.compile(
        r"<(h[1-4]|p|blockquote|li|pre)[^>]*>(.*?)</\1>",
        re.DOTALL,
    )
    parts: list[tuple[str, str]] = []
    for m in block_pattern.finditer(expanded):
        tag = m.group(1)
        inner = _strip_tags(m.group(2))
        if not inner:
            continue
        parts.append((tag, inner))

    if not parts:
        plain = _strip_tags(html)
        if plain:
            doc.add_paragraph(plain)
        return

    for tag, text in parts:
        if tag.startswith("h") and len(tag) == 2:
            level = int(tag[1])
            doc.add_heading(text, level=min(level + 1, 4))
        elif tag == "blockquote":
            doc.add_paragraph(text, style="Quote")
        elif tag == "li":
            doc.add_paragraph(text, style="List Bullet")
        elif tag == "pre":
            code_text = text.replace("\\n", "\n")
            code_para = doc.add_paragraph(code_text)
            for run in code_para.runs:
                run.font.name = "Courier New"
        else:
            doc.add_paragraph(text)


# =============================================================================
# DOCX Export
# =============================================================================


def create_docx_proposal(
    proposal: Proposal,
    sections: list[ProposalSection],
    rfp: RFP,
) -> bytes:
    """
    Generate a DOCX file from proposal sections.
    """
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx not installed. Run: pip install python-docx",
        )

    doc = Document()

    # Set document properties
    doc.core_properties.title = proposal.title
    doc.core_properties.author = "RFP Sniper"
    doc.core_properties.created = datetime.utcnow()

    # Title page
    title = doc.add_heading(proposal.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # RFP details
    details = doc.add_paragraph()
    details.add_run("Solicitation Number: ").bold = True
    details.add_run(rfp.solicitation_number)
    details.add_run("\n")
    details.add_run("Agency: ").bold = True
    details.add_run(rfp.agency)
    details.add_run("\n")
    if rfp.response_deadline:
        details.add_run("Due Date: ").bold = True
        details.add_run(rfp.response_deadline.strftime("%B %d, %Y"))

    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading("Table of Contents", 1)
    doc.add_paragraph("(Update field to generate TOC)")
    doc.add_page_break()

    # Executive Summary (if exists)
    if proposal.executive_summary:
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(proposal.executive_summary)
        doc.add_page_break()

    # Proposal sections
    for section in sorted(sections, key=lambda s: s.display_order):
        # Section heading
        doc.add_heading(f"{section.section_number}. {section.title}", 2)

        # Requirement text (if any)
        if section.requirement_text:
            req_para = doc.add_paragraph()
            req_para.add_run("Requirement: ").bold = True
            req_para.add_run(section.requirement_text[:500])
            if len(section.requirement_text) > 500:
                req_para.add_run("...")
            doc.add_paragraph()

        # Section content
        content = section.final_content
        if not content and section.generated_content:
            # Use clean text from generated content
            content = section.generated_content.get("clean_text", "")

        if content:
            _render_html_to_docx(doc, content)
        else:
            doc.add_paragraph("[Section content pending]")

        doc.add_paragraph()

    # Footer with generation info
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.add_run("Generated by RFP Sniper").italic = True
    footer.add_run(f" on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


@router.get("/proposals/{proposal_id}/docx")
async def export_proposal_docx(
    proposal_id: int,
    request: Request,
    current_user: UserAuth = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """
    Export a proposal to Microsoft Word (DOCX) format.
    """
    # Get proposal
    result = await session.execute(
        select(Proposal).where(
            Proposal.id == proposal_id,
            Proposal.user_id == current_user.id,
        )
    )
    proposal = result.scalar_one_or_none()

    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    # Policy check
    user_role = await get_user_policy_role(current_user.id, session)
    policy = evaluate(PolicyAction.EXPORT, proposal.classification, user_role)
    await log_audit_event(
        session,
        user_id=current_user.id,
        entity_type="proposal",
        entity_id=proposal_id,
        action="export_docx",
        metadata=policy.to_audit_dict(),
    )
    await _enforce_export_policy(
        policy=policy,
        current_user=current_user,
        session=session,
        request=request,
    )
    # Get sections
    sections_result = await session.execute(
        select(ProposalSection)
        .where(ProposalSection.proposal_id == proposal_id)
        .order_by(ProposalSection.display_order)
    )
    sections = list(sections_result.scalars().all())

    # Get RFP
    rfp_result = await session.execute(select(RFP).where(RFP.id == proposal.rfp_id))
    rfp = rfp_result.scalar_one_or_none()

    if not rfp:
        raise HTTPException(status_code=404, detail="Associated RFP not found")

    # Generate DOCX
    try:
        docx_bytes = create_docx_proposal(proposal, sections, rfp)
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

    await session.commit()

    # Create filename
    safe_title = "".join(c for c in proposal.title[:50] if c.isalnum() or c in " -_")
    filename = f"{safe_title}_{datetime.utcnow().strftime('%Y%m%d')}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# =============================================================================
# PDF Export
# =============================================================================


def create_pdf_proposal(
    proposal: Proposal,
    sections: list[ProposalSection],
    rfp: RFP,
) -> bytes:
    """
    Generate a PDF file from proposal sections.
    Uses weasyprint for HTML to PDF conversion.
    """
    try:
        from weasyprint import CSS, HTML
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="weasyprint not installed. Run: pip install weasyprint",
        )

    # Build HTML content
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                size: letter;
                margin: 1in;
                @bottom-center {{
                    content: "Page " counter(page) " of " counter(pages);
                    font-size: 10pt;
                    color: #666;
                }}
            }}
            body {{
                font-family: 'Georgia', serif;
                font-size: 12pt;
                line-height: 1.6;
                color: #333;
            }}
            h1 {{
                font-size: 24pt;
                text-align: center;
                margin-bottom: 0.5in;
                color: #1a1a1a;
            }}
            h2 {{
                font-size: 16pt;
                border-bottom: 2px solid #333;
                padding-bottom: 5pt;
                margin-top: 0.5in;
            }}
            .title-page {{
                text-align: center;
                page-break-after: always;
            }}
            .details {{
                margin: 1in auto;
                width: 80%;
            }}
            .details p {{
                margin: 0.2in 0;
            }}
            .requirement {{
                background: #f5f5f5;
                padding: 10pt;
                border-left: 3px solid #666;
                margin: 10pt 0;
                font-style: italic;
            }}
            .section-content {{
                margin: 0.2in 0;
            }}
            pre {{
                background: #f4f6fb;
                border: 1px solid #d7deed;
                border-radius: 6px;
                padding: 10pt;
                font-family: 'Courier New', monospace;
                font-size: 10pt;
                white-space: pre-wrap;
            }}
            code {{
                font-family: 'Courier New', monospace;
            }}
            .footer {{
                margin-top: 1in;
                font-size: 10pt;
                color: #666;
                font-style: italic;
            }}
        </style>
    </head>
    <body>
        <div class="title-page">
            <h1>{proposal.title}</h1>
            <div class="details">
                <p><strong>Solicitation Number:</strong> {rfp.solicitation_number}</p>
                <p><strong>Agency:</strong> {rfp.agency}</p>
                <p><strong>Due Date:</strong> {rfp.response_deadline.strftime("%B %d, %Y") if rfp.response_deadline else "TBD"}</p>
            </div>
        </div>
    """

    # Executive Summary
    if proposal.executive_summary:
        html_content += f"""
        <h2>Executive Summary</h2>
        <div class="section-content">
            {proposal.executive_summary.replace(chr(10), "<br>")}
        </div>
        """

    # Sections
    for section in sorted(sections, key=lambda s: s.display_order):
        html_content += f"""
        <h2>{section.section_number}. {section.title}</h2>
        """

        if section.requirement_text:
            html_content += f"""
            <div class="requirement">
                <strong>Requirement:</strong> {section.requirement_text[:500]}{"..." if len(section.requirement_text) > 500 else ""}
            </div>
            """

        content = section.final_content
        if not content and section.generated_content:
            content = section.generated_content.get("clean_text", "")

        if content:
            if _has_html(content):
                # TipTap HTML — inject directly
                html_content += f'<div class="section-content">{content}</div>'
            else:
                # Legacy plain text
                paragraphs = content.split("\n\n")
                for para in paragraphs:
                    if para.strip():
                        html_content += f"<p>{para.strip()}</p>"
        else:
            html_content += "<p><em>[Section content pending]</em></p>"

    # Footer
    html_content += f"""
        <div class="footer">
            Generated by RFP Sniper on {datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")}
        </div>
    </body>
    </html>
    """

    # Convert to PDF
    pdf_bytes = HTML(string=html_content).write_pdf()

    return pdf_bytes


@router.get("/proposals/{proposal_id}/pdf")
async def export_proposal_pdf(
    proposal_id: int,
    request: Request,
    current_user: UserAuth = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """
    Export a proposal to PDF format.
    """
    # Get proposal
    result = await session.execute(
        select(Proposal).where(
            Proposal.id == proposal_id,
            Proposal.user_id == current_user.id,
        )
    )
    proposal = result.scalar_one_or_none()

    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    # Policy check
    user_role = await get_user_policy_role(current_user.id, session)
    policy = evaluate(PolicyAction.EXPORT, proposal.classification, user_role)
    await log_audit_event(
        session,
        user_id=current_user.id,
        entity_type="proposal",
        entity_id=proposal_id,
        action="export_pdf",
        metadata=policy.to_audit_dict(),
    )
    await _enforce_export_policy(
        policy=policy,
        current_user=current_user,
        session=session,
        request=request,
    )

    # Get sections
    sections_result = await session.execute(
        select(ProposalSection)
        .where(ProposalSection.proposal_id == proposal_id)
        .order_by(ProposalSection.display_order)
    )
    sections = list(sections_result.scalars().all())

    # Get RFP
    rfp_result = await session.execute(select(RFP).where(RFP.id == proposal.rfp_id))
    rfp = rfp_result.scalar_one_or_none()

    if not rfp:
        raise HTTPException(status_code=404, detail="Associated RFP not found")

    # Generate PDF
    try:
        pdf_bytes = create_pdf_proposal(proposal, sections, rfp)
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

    await session.commit()

    # Create filename
    safe_title = "".join(c for c in proposal.title[:50] if c.isalnum() or c in " -_")
    filename = f"{safe_title}_{datetime.utcnow().strftime('%Y%m%d')}.pdf"

    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# =============================================================================
# Compliance Matrix Export
# =============================================================================


def _create_compliance_matrix_xlsx_bytes(matrix: ComplianceMatrix) -> bytes:
    """Render compliance matrix workbook bytes."""
    try:
        import openpyxl
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="openpyxl not installed. Run: pip install openpyxl",
        )

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Compliance Matrix"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    mandatory_fill = PatternFill(start_color="FFCCCB", end_color="FFCCCB", fill_type="solid")
    addressed_fill = PatternFill(start_color="90EE90", end_color="90EE90", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    headers = ["ID", "Section", "Requirement", "Type", "Importance", "Addressed", "Notes"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

    for row_num, req in enumerate(matrix.requirements, 2):
        ws.cell(row=row_num, column=1, value=req.get("id", ""))
        ws.cell(row=row_num, column=2, value=req.get("section", ""))
        ws.cell(row=row_num, column=3, value=req.get("requirement_text", ""))
        ws.cell(row=row_num, column=4, value=req.get("category", ""))
        ws.cell(row=row_num, column=5, value=req.get("importance", ""))
        ws.cell(row=row_num, column=6, value="Yes" if req.get("is_addressed") else "No")
        ws.cell(row=row_num, column=7, value=req.get("notes", ""))

        for col in range(1, 8):
            cell = ws.cell(row=row_num, column=col)
            cell.border = thin_border
            cell.alignment = Alignment(wrap_text=True, vertical="top")

        if req.get("importance") == "mandatory" and not req.get("is_addressed"):
            for col in range(1, 8):
                ws.cell(row=row_num, column=col).fill = mandatory_fill
        elif req.get("is_addressed"):
            ws.cell(row=row_num, column=6).fill = addressed_fill

    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 60
    ws.column_dimensions["D"].width = 15
    ws.column_dimensions["E"].width = 12
    ws.column_dimensions["F"].width = 10
    ws.column_dimensions["G"].width = 30

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()


def _build_compliance_package_manifest(
    proposal: Proposal,
    rfp: RFP,
    sections: list[ProposalSection],
    matrix: ComplianceMatrix | None,
) -> dict:
    requirements = matrix.requirements if matrix else []
    mandatory = [req for req in requirements if req.get("importance") == "mandatory"]
    addressed = [req for req in requirements if req.get("is_addressed")]
    return {
        "generated_at": datetime.utcnow().isoformat(),
        "proposal": {
            "id": proposal.id,
            "title": proposal.title,
            "status": proposal.status,
            "classification": proposal.classification,
        },
        "rfp": {
            "id": rfp.id,
            "title": rfp.title,
            "solicitation_number": rfp.solicitation_number,
            "agency": rfp.agency,
        },
        "summary": {
            "total_sections": len(sections),
            "completed_sections": sum(
                1 for section in sections if section.status.value in {"generated", "approved"}
            ),
            "total_requirements": len(requirements),
            "mandatory_requirements": len(mandatory),
            "addressed_requirements": len(addressed),
            "open_requirements": max(0, len(requirements) - len(addressed)),
        },
    }


def _enum_or_value(value: Any) -> Any:
    return value.value if hasattr(value, "value") else value


def _apply_cui_redaction_to_compliance_artifacts(
    *,
    source_trace_records: list[dict[str, Any]],
    section_records: list[dict[str, Any]],
    review_packets: list[dict[str, Any]],
) -> None:
    """Apply deterministic redaction transforms for CUI export artifacts."""
    for record in source_trace_records:
        record["document_title"] = "[REDACTED]"
        record["document_filename"] = "[REDACTED]"
        record["citation"] = "[REDACTED]" if record.get("citation") else None
        record["notes"] = "[REDACTED]" if record.get("notes") else None

    for section in section_records:
        if section.get("requirement_id"):
            section["requirement_id"] = "[REDACTED]"

    for packet in review_packets:
        if packet.get("proposal_title"):
            packet["proposal_title"] = "[REDACTED]"
        action_queue = packet.get("action_queue")
        if isinstance(action_queue, list):
            for item in action_queue:
                if not isinstance(item, dict):
                    continue
                if item.get("recommended_action"):
                    item["recommended_action"] = "[REDACTED]"
                if item.get("rationale"):
                    item["rationale"] = "[REDACTED]"
        exit_criteria = packet.get("recommended_exit_criteria")
        if isinstance(exit_criteria, list):
            packet["recommended_exit_criteria"] = [
                "[REDACTED]" if criterion else criterion for criterion in exit_criteria
            ]


@router.get("/rfps/{rfp_id}/compliance-matrix/xlsx")
async def export_compliance_matrix_xlsx(
    rfp_id: int,
    current_user: UserAuth = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """
    Export compliance matrix to Excel (XLSX) format.
    """
    # Get RFP
    result = await session.execute(
        select(RFP).where(
            RFP.id == rfp_id,
            RFP.user_id == current_user.id,
        )
    )
    rfp = result.scalar_one_or_none()

    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")

    # Get compliance matrix
    matrix_result = await session.execute(
        select(ComplianceMatrix).where(ComplianceMatrix.rfp_id == rfp_id)
    )
    matrix = matrix_result.scalar_one_or_none()

    if not matrix:
        raise HTTPException(status_code=404, detail="Compliance matrix not found")

    matrix_bytes = _create_compliance_matrix_xlsx_bytes(matrix)

    filename = (
        f"compliance_matrix_{rfp.solicitation_number}_{datetime.utcnow().strftime('%Y%m%d')}.xlsx"
    )

    return StreamingResponse(
        io.BytesIO(matrix_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/proposals/{proposal_id}/compliance-package/zip")
async def export_proposal_compliance_package(
    proposal_id: int,
    request: Request,
    current_user: UserAuth = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """
    Export an evaluator-ready compliance package ZIP.
    Includes proposal DOCX, compliance matrix XLSX (if available), section evidence summary,
    and a manifest describing package coverage.
    """
    proposal = (
        await session.execute(
            select(Proposal).where(
                Proposal.id == proposal_id,
                Proposal.user_id == current_user.id,
            )
        )
    ).scalar_one_or_none()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    user_role = await get_user_policy_role(current_user.id, session)
    policy = evaluate(PolicyAction.EXPORT, proposal.classification, user_role)
    await log_audit_event(
        session,
        user_id=current_user.id,
        entity_type="proposal",
        entity_id=proposal_id,
        action="export_compliance_package",
        metadata=policy.to_audit_dict(),
    )
    await _enforce_export_policy(
        policy=policy,
        current_user=current_user,
        session=session,
        request=request,
    )
    org_security_policy = await get_user_org_security_policy(current_user.id, session)
    proposal_classification = str(_enum_or_value(proposal.classification)).lower()
    is_cui_export = proposal_classification == "cui"
    watermark_applied = bool(
        is_cui_export and org_security_policy.get("apply_cui_watermark_to_sensitive_exports", True)
    )
    redaction_applied = bool(
        is_cui_export and org_security_policy.get("apply_cui_redaction_to_sensitive_exports", False)
    )

    sections = list(
        (
            await session.execute(
                select(ProposalSection)
                .where(ProposalSection.proposal_id == proposal_id)
                .order_by(ProposalSection.display_order)
            )
        )
        .scalars()
        .all()
    )
    rfp = (await session.execute(select(RFP).where(RFP.id == proposal.rfp_id))).scalar_one_or_none()
    if not rfp:
        raise HTTPException(status_code=404, detail="Associated RFP not found")

    matrix = (
        await session.execute(select(ComplianceMatrix).where(ComplianceMatrix.rfp_id == rfp.id))
    ).scalar_one_or_none()

    docx_bytes = create_docx_proposal(proposal, sections, rfp)
    matrix_bytes = _create_compliance_matrix_xlsx_bytes(matrix) if matrix else None
    section_ids = [section.id for section in sections if section.id is not None]

    section_records = [
        {
            "id": section.id,
            "section_number": section.section_number,
            "title": section.title,
            "status": section.status.value,
            "requirement_id": section.requirement_id,
            "word_count": section.word_count,
            "quality_score": section.quality_score,
            "has_content": bool(
                section.final_content
                or (section.generated_content and section.generated_content.get("clean_text"))
            ),
        }
        for section in sections
    ]

    source_trace_records: list[dict[str, Any]] = []
    section_evidence_count: dict[int, int] = defaultdict(int)
    if section_ids:
        section_evidence = list(
            (
                await session.execute(
                    select(SectionEvidence).where(SectionEvidence.section_id.in_(section_ids))
                )
            )
            .scalars()
            .all()
        )
        document_ids = sorted(
            {row.document_id for row in section_evidence if row.document_id is not None}
        )
        chunk_ids = sorted({row.chunk_id for row in section_evidence if row.chunk_id is not None})

        documents_by_id: dict[int, KnowledgeBaseDocument] = {}
        chunks_by_id: dict[int, DocumentChunk] = {}

        if document_ids:
            documents = list(
                (
                    await session.execute(
                        select(KnowledgeBaseDocument).where(
                            KnowledgeBaseDocument.id.in_(document_ids)
                        )
                    )
                )
                .scalars()
                .all()
            )
            documents_by_id = {doc.id: doc for doc in documents if doc.id is not None}

        if chunk_ids:
            chunks = list(
                (
                    await session.execute(
                        select(DocumentChunk).where(DocumentChunk.id.in_(chunk_ids))
                    )
                )
                .scalars()
                .all()
            )
            chunks_by_id = {chunk.id: chunk for chunk in chunks if chunk.id is not None}

        sections_by_id = {section.id: section for section in sections if section.id is not None}

        for evidence in section_evidence:
            section = sections_by_id.get(evidence.section_id)
            document = documents_by_id.get(evidence.document_id)
            chunk = chunks_by_id.get(evidence.chunk_id) if evidence.chunk_id else None
            section_evidence_count[evidence.section_id] += 1
            source_trace_records.append(
                {
                    "evidence_id": evidence.id,
                    "section_id": evidence.section_id,
                    "section_number": section.section_number if section else None,
                    "section_title": section.title if section else None,
                    "document_id": evidence.document_id,
                    "document_title": document.title if document else None,
                    "document_filename": document.original_filename if document else None,
                    "document_type": (_enum_or_value(document.document_type) if document else None),
                    "chunk_id": evidence.chunk_id,
                    "chunk_page_number": chunk.page_number if chunk else None,
                    "citation": evidence.citation,
                    "notes": evidence.notes,
                    "linked_at": evidence.created_at.isoformat(),
                }
            )

    reviews = list(
        (
            await session.execute(
                select(ProposalReview)
                .where(ProposalReview.proposal_id == proposal_id)
                .order_by(ProposalReview.created_at.desc())
            )
        )
        .scalars()
        .all()
    )

    review_comments_by_section: dict[int, dict[str, Any]] = defaultdict(
        lambda: {
            "total_comments": 0,
            "open_comments": 0,
            "resolved_comments": 0,
            "open_critical": 0,
            "open_major": 0,
            "last_comment_at": None,
        }
    )
    review_packets: list[dict[str, Any]] = []
    review_outcomes: list[dict[str, Any]] = []
    if reviews:
        from app.api.routes.reviews import get_review_packet

        review_ids = [review.id for review in reviews if review.id is not None]
        if review_ids:
            comments = list(
                (
                    await session.execute(
                        select(ReviewComment).where(ReviewComment.review_id.in_(review_ids))
                    )
                )
                .scalars()
                .all()
            )
            actionable_statuses = {"open", "assigned", "addressed"}
            resolved_statuses = {"closed", "verified", "rejected"}
            for comment in comments:
                if comment.section_id is None:
                    continue
                status = str(_enum_or_value(comment.status))
                severity = str(_enum_or_value(comment.severity))
                section_stats = review_comments_by_section[comment.section_id]
                section_stats["total_comments"] += 1
                if status in actionable_statuses:
                    section_stats["open_comments"] += 1
                    if severity == "critical":
                        section_stats["open_critical"] += 1
                    if severity == "major":
                        section_stats["open_major"] += 1
                if status in resolved_statuses:
                    section_stats["resolved_comments"] += 1
                if (
                    section_stats["last_comment_at"] is None
                    or comment.created_at > section_stats["last_comment_at"]
                ):
                    section_stats["last_comment_at"] = comment.created_at

        for review in reviews:
            packet = await get_review_packet(
                review_id=review.id,  # type: ignore[arg-type]
                current_user=current_user,
                session=session,
            )
            packet_payload = packet.model_dump(mode="json")
            review_packets.append(packet_payload)
            review_outcomes.append(
                {
                    "review_id": packet_payload["review_id"],
                    "review_type": packet_payload["review_type"],
                    "review_status": packet_payload["review_status"],
                    "overall_risk_level": packet_payload["risk_summary"]["overall_risk_level"],
                    "open_critical": packet_payload["risk_summary"]["open_critical"],
                    "open_major": packet_payload["risk_summary"]["open_major"],
                    "checklist_pass_rate": packet_payload["checklist_summary"]["pass_rate"],
                }
            )

    section_decision_records = []
    for section in sections:
        stats = review_comments_by_section.get(section.id)
        last_comment_at = None
        if stats and stats.get("last_comment_at"):
            last_comment_at = stats["last_comment_at"].isoformat()
        section_decision_records.append(
            {
                "section_id": section.id,
                "section_number": section.section_number,
                "title": section.title,
                "status": _enum_or_value(section.status),
                "quality_score": section.quality_score,
                "requirement_id": section.requirement_id,
                "word_count": section.word_count,
                "linked_evidence_count": section_evidence_count.get(section.id, 0),
                "review_outcomes": {
                    "total_comments": stats["total_comments"] if stats else 0,
                    "open_comments": stats["open_comments"] if stats else 0,
                    "resolved_comments": stats["resolved_comments"] if stats else 0,
                    "open_critical": stats["open_critical"] if stats else 0,
                    "open_major": stats["open_major"] if stats else 0,
                    "last_comment_at": last_comment_at,
                },
            }
        )

    bid_stress_test_payload: dict[str, Any] | None = None
    scorecards = list(
        (
            await session.execute(
                select(BidScorecard)
                .where(BidScorecard.rfp_id == rfp.id, BidScorecard.user_id == current_user.id)
                .order_by(BidScorecard.created_at.desc())
            )
        )
        .scalars()
        .all()
    )
    if scorecards:
        from app.services.bid_scenario_service import BidScenarioService

        baseline_scorecard = next(
            (card for card in scorecards if card.criteria_scores), scorecards[0]
        )
        simulation = BidScenarioService().simulate(scorecard=baseline_scorecard)
        bid_stress_test_payload = {
            "generated_at": datetime.utcnow().isoformat(),
            "baseline_scorecard_id": baseline_scorecard.id,
            **simulation,
        }

    if redaction_applied:
        _apply_cui_redaction_to_compliance_artifacts(
            source_trace_records=source_trace_records,
            section_records=section_records,
            review_packets=review_packets,
        )

    manifest = _build_compliance_package_manifest(proposal, rfp, sections, matrix)
    manifest["summary"]["source_trace_links"] = len(source_trace_records)
    manifest["summary"]["section_decisions"] = len(section_decision_records)
    manifest["summary"]["review_packets_included"] = len(review_packets)
    manifest["summary"]["bid_stress_test_included"] = bool(bid_stress_test_payload)
    manifest["artifacts"] = {
        "proposal_docx": True,
        "compliance_matrix_xlsx": bool(matrix_bytes),
        "source_trace": bool(source_trace_records),
        "section_decisions": bool(section_decision_records),
        "review_packets": bool(review_packets),
        "bid_stress_test": bool(bid_stress_test_payload),
        "classification_watermark": watermark_applied,
        "cui_redaction": redaction_applied,
    }
    manifest["policy_actions"] = {
        "watermark_applied": watermark_applied,
        "redaction_applied": redaction_applied,
        "watermark_targets": ["classification-watermark.txt"] if watermark_applied else [],
        "redaction_targets": (
            [
                "source-trace.json",
                "sections.json",
                "reviews/review-packets.json",
            ]
            if redaction_applied
            else []
        ),
    }

    package_bytes = io.BytesIO()
    with zipfile.ZipFile(package_bytes, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("manifest.json", json.dumps(manifest, indent=2))
        archive.writestr("sections.json", json.dumps(section_records, indent=2))
        archive.writestr("source-trace.json", json.dumps(source_trace_records, indent=2))
        archive.writestr("section-decisions.json", json.dumps(section_decision_records, indent=2))
        archive.writestr("reviews/review-packets.json", json.dumps(review_packets, indent=2))
        archive.writestr("reviews/review-outcomes.json", json.dumps(review_outcomes, indent=2))
        if bid_stress_test_payload:
            archive.writestr(
                "capture/bid-stress-test.json",
                json.dumps(bid_stress_test_payload, indent=2),
            )
        archive.writestr("proposal.docx", docx_bytes)
        if matrix_bytes:
            archive.writestr("compliance-matrix.xlsx", matrix_bytes)
        if watermark_applied:
            archive.writestr(
                "classification-watermark.txt",
                (
                    "CUI // CONTROLLED UNCLASSIFIED INFORMATION\n"
                    "This package contains data governed by CUI handling controls.\n"
                    "Distribution and retention must follow organization policy.\n"
                ),
            )
        archive.writestr(
            "README.txt",
            (
                "Compliance package generated by GovTech Sniper.\n"
                "- proposal.docx: proposal narrative export\n"
                "- compliance-matrix.xlsx: requirement traceability sheet\n"
                "- sections.json: section-level status and quality metadata\n"
                "- source-trace.json: section evidence links to source documents/chunks\n"
                "- section-decisions.json: section decisions with review outcome rollups\n"
                "- reviews/review-packets.json: risk-ranked pink/red/gold review packets\n"
                "- reviews/review-outcomes.json: condensed review outcome metrics\n"
                "- capture/bid-stress-test.json: FAR/Section M scenario simulation output\n"
                "- manifest.json: package summary and coverage counts\n"
                "- classification-watermark.txt: CUI package handling notice (when enabled)\n"
                "- Redaction mode (CUI policy): source trace, sections, and review packet text are redacted\n"
            ),
        )

    await session.commit()
    package_bytes.seek(0)
    safe_title = "".join(c for c in proposal.title[:50] if c.isalnum() or c in " -_").strip()
    filename = (
        f"{safe_title or 'proposal'}_compliance_package_{datetime.utcnow().strftime('%Y%m%d')}.zip"
    )
    return StreamingResponse(
        package_bytes,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
