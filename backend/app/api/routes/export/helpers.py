"""
Export Routes - Shared Helpers
==============================
HTML-to-DOCX rendering, policy enforcement, and utility functions.
"""

from __future__ import annotations

import html
import re
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from docx.document import Document

import structlog
from fastapi import HTTPException, Request
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import (
    STEP_UP_REQUIRED_HEADER,
    UserAuth,
    get_step_up_code,
    get_user_org_security_policy,
    verify_step_up_code,
)
from app.services.audit_service import log_audit_event
from app.services.policy_engine import PolicyDecision, PolicyResult

logger = structlog.get_logger(__name__)

_TAG_RE = re.compile(r"<[^>]+>")


async def enforce_export_policy(
    *,
    policy: PolicyResult,
    current_user: UserAuth,
    session: AsyncSession,
    request: Request,
    step_up_code: str | None = None,
) -> None:
    if policy.decision == PolicyDecision.DENY:
        await session.commit()
        raise HTTPException(status_code=403, detail=policy.reason)
    if policy.decision != PolicyDecision.STEP_UP:
        return

    org_policy = await get_user_org_security_policy(current_user.id, session)
    if not org_policy.get("require_step_up_for_sensitive_exports", True):
        return

    supplied_code = get_step_up_code(request, step_up_code)
    if await verify_step_up_code(current_user.id, session, supplied_code):
        await log_audit_event(
            session,
            user_id=current_user.id,
            entity_type="security",
            entity_id=current_user.id,
            action="security.step_up.challenge_succeeded",
            metadata={
                "channel": "export",
                **policy.to_audit_dict(),
            },
        )
        await session.commit()
        return

    await log_audit_event(
        session,
        user_id=current_user.id,
        entity_type="security",
        entity_id=current_user.id,
        action="security.step_up.challenge_failed",
        metadata={
            "channel": "export",
            **policy.to_audit_dict(),
        },
    )
    await session.commit()
    raise HTTPException(
        status_code=403,
        detail=policy.reason,
        headers={STEP_UP_REQUIRED_HEADER: "true"},
    )


def strip_tags(content: str) -> str:
    """Remove all HTML tags, returning plain text."""
    return html.unescape(_TAG_RE.sub("", content)).strip()


def has_html(text: str) -> bool:
    return bool(re.search(r"<[a-z][\s\S]*>", text, re.IGNORECASE))


def render_html_to_docx(doc: Document, html_str: str) -> None:
    """Convert HTML content from TipTap into python-docx paragraphs."""
    if not has_html(html_str):
        for para in html_str.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        return

    def _expand_lists(h: str) -> str:
        return re.sub(
            r"<(ul|ol)[^>]*>(.*?)</\1>",
            lambda m: m.group(2),
            h,
            flags=re.DOTALL,
        )

    expanded = _expand_lists(html_str)

    block_pattern = re.compile(
        r"<(h[1-4]|p|blockquote|li|pre)[^>]*>(.*?)</\1>",
        re.DOTALL,
    )
    parts: list[tuple[str, str]] = []
    for m in block_pattern.finditer(expanded):
        tag = m.group(1)
        inner = strip_tags(m.group(2))
        if not inner:
            continue
        parts.append((tag, inner))

    if not parts:
        plain = strip_tags(html_str)
        if plain:
            doc.add_paragraph(plain)
        return

    for tag, text in parts:
        if tag.startswith("h") and len(tag) == 2:
            level = int(tag[1])
            doc.add_heading(text, level=min(level + 1, 4))
        elif tag == "blockquote":
            doc.add_paragraph(text, style="Quote")
        elif tag == "li":
            doc.add_paragraph(text, style="List Bullet")
        elif tag == "pre":
            code_text = text.replace("\\n", "\n")
            code_para = doc.add_paragraph(code_text)
            for run in code_para.runs:
                run.font.name = "Courier New"
        else:
            doc.add_paragraph(text)


def enum_or_value(value: Any) -> Any:
    return value.value if hasattr(value, "value") else value


def apply_cui_redaction_to_compliance_artifacts(
    *,
    source_trace_records: list[dict[str, Any]],
    section_records: list[dict[str, Any]],
    review_packets: list[dict[str, Any]],
) -> None:
    """Apply deterministic redaction transforms for CUI export artifacts."""
    for record in source_trace_records:
        record["document_title"] = "[REDACTED]"
        record["document_filename"] = "[REDACTED]"
        record["citation"] = "[REDACTED]" if record.get("citation") else None
        record["notes"] = "[REDACTED]" if record.get("notes") else None

    for section in section_records:
        if section.get("requirement_id"):
            section["requirement_id"] = "[REDACTED]"

    for packet in review_packets:
        if packet.get("proposal_title"):
            packet["proposal_title"] = "[REDACTED]"
        action_queue = packet.get("action_queue")
        if isinstance(action_queue, list):
            for item in action_queue:
                if not isinstance(item, dict):
                    continue
                if item.get("recommended_action"):
                    item["recommended_action"] = "[REDACTED]"
                if item.get("rationale"):
                    item["rationale"] = "[REDACTED]"
        exit_criteria = packet.get("recommended_exit_criteria")
        if isinstance(exit_criteria, list):
            packet["recommended_exit_criteria"] = [
                "[REDACTED]" if criterion else criterion for criterion in exit_criteria
            ]
